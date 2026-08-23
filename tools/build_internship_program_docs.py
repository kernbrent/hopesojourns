from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "admin" / "internship-program"
LOGO = ROOT / "assets" / "hope-sojourns-logo.png"

FOREST = "174F43"
FOREST_DARK = "0F352D"
GOLD = "D99B42"
CORAL = "C9674E"
CREAM = "FFF8E9"
MIST = "E8F1ED"
PALE_GOLD = "FFF3DD"
INK = "19322B"
MUTED = "5A6E67"
LINE = "D7E2DD"
WHITE = "FFFFFF"
RED = "9B2F2F"

BODY_FONT = "Calibri"
TITLE_FONT = "Georgia"
TODAY_LABEL = "August 2026"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=BODY_FONT, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(index, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_border(paragraph, color=GOLD, size=18, space=6):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_link(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), FOREST)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc, compact=False):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(10.5 if compact else 11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6 if compact else 8)
    normal.paragraph_format.line_spacing = 1.18 if compact else 1.25

    title = doc.styles["Title"]
    title.font.name = TITLE_FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), TITLE_FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), TITLE_FONT)
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(FOREST_DARK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    heading_tokens = {
        "Heading 1": (18, FOREST_DARK, 18, 9),
        "Heading 2": (14, FOREST, 14, 7),
        "Heading 3": (11.5, CORAL, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = TITLE_FONT if name != "Heading 3" else BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(10.5 if compact else 11)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.42 if compact else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.2 if compact else -0.25)
        style.paragraph_format.space_after = Pt(4 if compact else 6)
        style.paragraph_format.line_spacing = 1.18 if compact else 1.22


def add_running_furniture(doc, label):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"HOPE SOJOURNS  |  {label.upper()}")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(p, color=LINE, size=8, space=5)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Working draft  |  Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)


def new_document(title, subtitle, label, status="Working draft", compact=False):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    configure_styles(doc, compact=compact)
    add_running_furniture(doc, label)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(LOGO), width=Inches(2.15))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label.upper())
    set_run_font(run, size=9, color=GOLD, bold=True)

    p = doc.add_paragraph(style="Title")
    p.add_run(title)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=12, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(f"{status}  |  {TODAY_LABEL}")
    set_run_font(run, size=9, color=FOREST, bold=True)
    set_paragraph_border(p, color=GOLD, size=16, space=6)

    doc.core_properties.title = title
    doc.core_properties.subject = "Hope Sojourns college ministry internship program"
    doc.core_properties.author = "Hope Sojourns"
    doc.core_properties.keywords = "Hope Sojourns, internship, ministry, university, student support"
    return doc


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text="", bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.extend([start, num_fmt, level_text, level_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.22
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        run = p.add_run(item)
        set_run_font(run)


def add_callout(doc, title, text, fill=PALE_GOLD, accent=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    set_run_font(run, size=8.5, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK, bold=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=FOREST, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=WHITE, bold=True)
    for row_values in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, "F7FAF8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_labeled_fields(doc, fields):
    for label, placeholder in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        lead = p.add_run(f"{label}: ")
        set_run_font(lead, size=10.5, bold=True, color=FOREST_DARK)
        value = p.add_run(f"[{placeholder}]")
        set_run_font(value, size=10.5, color=MUTED, italic=True)


def add_checkbox_list(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.space_after = Pt(5)
        check = p.add_run("☐  ")
        set_run_font(check, size=11, color=FOREST)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_signature_table(doc, signers):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4680, 4680])
    for signer in signers:
        row = table.add_row()
        set_row_cant_split(row)
        left, right = row.cells
        set_cell_shading(left, "FBFCFB")
        set_cell_shading(right, "FBFCFB")
        left.text = f"{signer}\nSignature: ______________________________\nPrinted name: ___________________________"
        right.text = "Date: __________________\nTitle/organization: _____________________\nEmail: _________________________________"
        for cell in (left, right):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    set_run_font(run, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_legal_review_note(doc):
    add_callout(
        doc,
        "Review before use",
        "This is an operational draft, not legal advice. Hope Sojourns should have Texas counsel, its insurance broker, and each participating university review the document before adoption. Country-specific counsel may be needed for international placements.",
        fill="FFF0EC",
        accent=RED,
    )


def save(doc, filename):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc.save(path)
    return path


def build_business_model():
    doc = new_document(
        "College Ministry Internship Program",
        "Business model, operating approach, and launch plan",
        "Program strategy",
        status="Internal working model",
        compact=False,
    )
    add_callout(
        doc,
        "Recommended position",
        "Hope Sojourns should operate as a vetted internship-placement and field-support partner: connecting students and universities with ministry hosts, preparing participants, helping them transition on site, mentoring them weekly, and maintaining accountable communication.",
    )

    add_heading(doc, "1. Executive summary", 1)
    add_paragraph(doc, "The program is feasible as a carefully limited pilot. Its strongest value is trust: credible ministry partners, academic alignment, thoughtful preparation, personal support, and clear communication before, during, and after the placement.")
    add_bullets(doc, [
        "Begin with four to eight students and no more than six active interns at one time.",
        "Launch with established ministry relationships and placements lasting primarily six to twelve weeks.",
        "Allow universities to refer or promote candidates, while Hope Sojourns and the host ministry jointly screen and accept each student.",
        "Use written agreements that separate university, Hope Sojourns, host ministry, and student responsibilities.",
        "Charge a transparent Program Support Fee and quote direct leader travel separately at actual cost.",
        "Create donor- and university-supported scholarships so participation is not limited to students with substantial resources.",
    ])

    add_heading(doc, "2. Program promise", 1)
    add_paragraph(doc, "Hope Sojourns connects college students with vetted ministry internships and surrounds them with academic coordination, practical preparation, on-site transition support, weekly mentoring, and accountable care throughout the experience.")
    add_heading(doc, "Who the program serves", 2)
    add_table(doc, ["Audience", "Value delivered"], [
        ("Students", "A meaningful placement, clear expectations, preparation, mentoring, advocacy, and a supported transition into a new community."),
        ("Universities", "Vetted opportunities, named supervision, learning-plan alignment, consistent communication, and documented evaluations."),
        ("Ministry partners", "Prepared students, well-defined responsibilities, one Hope Sojourns contact, and help resolving fit or performance concerns."),
        ("Families", "A clear emergency-contact process and confidence that the student has preparation, supervision, and ongoing support."),
    ], [1900, 7460])

    add_heading(doc, "3. Roles and accountability", 1)
    add_table(doc, ["Party", "Primary responsibilities"], [
        ("University or college", "Determines academic eligibility, credit, assignments, required hours, faculty oversight, and grades."),
        ("Hope Sojourns", "Vets the host, screens and matches the student, coordinates agreements, provides preparation, supports arrival, meets weekly, communicates with the university, and manages escalation."),
        ("Host ministry", "Provides a qualified on-site supervisor, meaningful work, regular feedback, safe working and living arrangements, local orientation, and required evaluations."),
        ("Student", "Completes academic and ministry expectations, follows policies, maintains communication, obtains required documents and insurance, and pays disclosed expenses."),
        ("Parent or emergency contact", "Receives information authorized by the adult student and participates when requested or when a qualifying emergency requires contact."),
    ], [1900, 7460])

    add_heading(doc, "4. Student experience", 1)
    add_numbered(doc, [
        "Explore: The student reviews current ministry possibilities and submits an interest form.",
        "Assess fit: Hope Sojourns learns about the student's field of study, maturity, goals, availability, and support needs.",
        "Confirm the placement: The host ministry interviews the student and the university confirms academic requirements when credit is involved.",
        "Prepare: Hope Sojourns leads three sessions covering role expectations, cultural humility, logistics, safety, and emergency communication.",
        "Arrive and settle: A Hope Sojourns leader accompanies the student when included, or coordinates a documented local arrival plan at established sites.",
        "Serve and reflect: The student receives daily host supervision and meets weekly with Hope Sojourns for mentoring, reflection, and problem-solving.",
        "Close well: The student, ministry, university, and Hope Sojourns complete evaluations, learning reflection, and next-step recommendations.",
    ])

    add_heading(doc, "5. Current placement portfolio", 1)
    add_paragraph(doc, "All opportunities remain subject to partner confirmation, university approval, appropriate supervision, background checks, travel requirements, and fit.")
    add_table(doc, ["Location", "Potential ministry", "Illustrative learning and service", "Length"], [
        ("Athens, Greece", "New Start Ministries", "Compassionate outreach and practical support for women affected by trafficking and sexual exploitation, with an emphasis on dignity, safety, and restoration.", "1-9 months"),
        ("Athens, Greece", "Glocal Cafe", "Help refugees navigate social and medical support and access food, clothing, and other practical resources through a publicly supported community organization.", "1-9 months"),
        ("Northern Arkansas", "Shephard of the Ozarks (SOTO)", "Support Christian camp programming and learn alongside staff caring for camp animals, wooded property, and outdoor ministry spaces.", "1-3 months"),
        ("Dallas, Texas", "Metro Relief of Dallas", "Learn alongside case workers, support grant development, and participate in relationship-centered outreach with people experiencing homelessness.", "1-3 months"),
        ("Mexico City, Mexico", "God's Kitchen", "Serve people experiencing homelessness, displaced families and refugees, and individuals rebuilding their lives after addiction.", "2-6 months"),
        ("England", "The Light Group", "A developing ministry placement whose final responsibilities will be defined with the partner and participating university.", "1-4 months"),
    ], [1500, 1900, 4460, 1500], font_size=8.6)

    add_heading(doc, "6. Preparation and support model", 1)
    add_heading(doc, "Pre-departure meetings", 2)
    add_table(doc, ["Session", "Purpose"], [
        ("1. Calling, goals, and expectations", "Clarify the student's learning goals, role, schedule, boundaries, university requirements, and what success looks like."),
        ("2. Culture, dignity, and ministry practice", "Prepare the student for cultural humility, trauma-aware interaction, ethical storytelling, safeguarding, confidentiality, and respectful service."),
        ("3. Travel, safety, and communication", "Review travel documents, insurance, local arrival, emergency contacts, technology, personal safety, and incident reporting."),
    ], [2700, 6660])
    add_heading(doc, "Weekly check-in rhythm", 2)
    add_bullets(doc, [
        "Well-being and adjustment",
        "Progress toward learning objectives",
        "Relationship with the on-site supervisor and team",
        "Cultural observations and ethical questions",
        "Concerns, incidents, or support needed",
        "Commitments and follow-up before the next meeting",
    ])
    add_callout(doc, "Scalability decision", "Accompanying every student is a strong launch-year differentiator but a long-term capacity constraint. Group students into common arrival windows and allow established sites to use a documented host-led arrival option when appropriate.", fill=MIST, accent=FOREST)

    add_heading(doc, "7. Financial model", 1)
    add_paragraph(doc, "The public-facing term should be Program Support Fee because it describes the value delivered more clearly than administration fee.")
    add_table(doc, ["Placement", "Recommended pilot fee"], [
        ("Domestic, 1-2 months", "$1,250"),
        ("Domestic, 2-4 months", "$1,750"),
        ("International, 1-2 months", "$2,750"),
        ("International, 3-4 months", "$3,500"),
        ("International, 5-6 months", "$4,500"),
        ("International, 7-9 months", "$5,750"),
    ], [6500, 2860])
    add_heading(doc, "What the Program Support Fee covers", 2)
    add_bullets(doc, [
        "Applicant assessment, matching, and host coordination",
        "University learning-agreement support",
        "Three structured preparation meetings",
        "Weekly individual check-ins and documentation",
        "University and host communication",
        "Midpoint and final reporting",
        "Arrival and transition planning",
        "Hope Sojourns leader professional time during launch travel",
        "Routine problem-solving and emergency coordination",
    ])
    add_heading(doc, "Student financial responsibility", 2)
    add_bullets(doc, [
        "University tuition and fees, when required by the student's institution",
        "The student's airfare and other travel",
        "Housing, food, and local transportation unless provided or subsidized by the local ministry",
        "Passport, visa, background check, vaccinations, and personal expenses",
        "Required travel health, medical evacuation, and trip-protection coverage",
        "The Program Support Fee",
        "A disclosed share of direct Hope Sojourns leader travel expenses when accompaniment is included",
    ])
    add_heading(doc, "Leader travel policy", 2)
    add_bullets(doc, [
        "Quote airfare, lodging, ground transportation, visas, and a published meal allowance before commitment.",
        "Bill direct travel expenses at actual cost without airfare markup.",
        "Split leader travel expenses equitably when multiple students share the same arrival trip.",
        "Use a written maximum authorized amount and document any material change before purchase.",
        "Keep leader professional time inside the Program Support Fee rather than charging a separate daily consulting rate.",
    ])
    add_callout(doc, "Access and equity", "Because students may also pay tuition and serve without wages, Hope Sojourns should seek ministry-provided housing, food, local transportation, or stipends whenever possible and establish a donor-supported Internship Access Fund.", fill="FFF0EC", accent=CORAL)

    add_heading(doc, "8. University partnership strategy", 1)
    add_paragraph(doc, "Begin with faculty internship coordinators and departments where current placements can support clear learning objectives. University-wide contracts can follow after the pilot produces credible outcomes.")
    add_heading(doc, "Priority academic areas", 2)
    add_bullets(doc, [
        "Ministry, theology, nonprofit management, and organizational leadership",
        "Communications, marketing, journalism, photography, and digital media",
        "Business, accounting, grant development, project management, and hospitality",
        "International studies, sociology, community development, and human services",
        "Education support where the host and university approve appropriate supervision",
    ])
    add_paragraph(doc, "Clinical social work, counseling, medical, and licensed education placements should wait until a host can satisfy the applicable supervision and accreditation requirements.")
    add_heading(doc, "First conversation", 2)
    add_numbered(doc, [
        "Present a one-page overview and two complete placement descriptions.",
        "Ask the department to explain its required hours, learning agreement, supervisor qualifications, evaluation schedule, and insurance requirements.",
        "Agree on how students will be referred, screened, approved, and supported.",
        "Pilot with one or two students and complete a joint review before expanding.",
    ])

    add_heading(doc, "9. Safeguards and operating standards", 1)
    add_bullets(doc, [
        "Signed university placement agreement, host ministry agreement, student participation agreement, and internship learning agreement",
        "Code of conduct and safeguarding/misconduct-reporting policy",
        "Background-check policy for placements involving children or vulnerable people",
        "Published cancellation and refund policy",
        "Trip-specific emergency plan and a backup Hope Sojourns emergency contact",
        "Written parent/emergency-contact authorization for adult students",
        "Data privacy, access, breach-response, and record-retention standards",
        "General liability, professional liability, travel medical, evacuation, and other coverage confirmed with a qualified broker",
        "Country-specific visa, work-permission, mandatory-reporting, and local-law review",
    ])
    add_paragraph(doc, "Routine parent communication should be based on the adult student's written authorization. Health or safety emergencies require the emergency plan and applicable law to control the response.")

    add_heading(doc, "10. Year-one pilot", 1)
    add_table(doc, ["Phase", "Actions", "Success evidence"], [
        ("Foundation | 0-45 days", "Finalize the program model, vet priority sites, confirm insurance, review templates with counsel, and interview university coordinators.", "Two placement-ready hosts and an approved operating packet."),
        ("Recruitment | 45-90 days", "Publish student-facing information, provide university briefs, accept applications, screen students, and confirm academic requirements.", "Four to eight qualified applicants and two university relationships."),
        ("Pilot delivery | 3-12 months", "Run preparation, accompany launch placements as promised, complete weekly check-ins, document incidents, and collect evaluations.", "At least 90% completion and strong student, host, and university feedback."),
        ("Review and scale | after pilot", "Compare actual hours and costs with fees, revise policies, strengthen scholarships, and group future arrivals by destination.", "Repeat host/university interest and a sustainable cost per placement."),
    ], [1500, 5060, 2800], font_size=8.5)

    add_heading(doc, "11. Performance measures", 1)
    add_bullets(doc, [
        "Application-to-placement conversion and reasons applicants do not proceed",
        "Student completion rate and learning-goal attainment",
        "Student, host supervisor, and university satisfaction",
        "Incidents, response time, resolution quality, and lessons learned",
        "Direct staff hours, direct cost, and contribution margin by placement",
        "Housing, food, transportation, stipend, and scholarship support secured",
        "Repeat placements from the same university and host ministry",
    ])

    add_heading(doc, "12. Source notes", 1)
    sources = [
        ("National Association of Colleges and Employers - Internship definition and quality standards", "https://naceweb.org/internships/"),
        ("U.S. Department of Labor - Fact Sheet #71 on internship programs", "https://www.dol.gov/agencies/whd/fact-sheets/71-flsa-internships"),
        ("Southern Methodist University - Internship Learning Contract", "https://www.smu.edu/-/media/site/dedman/studentresources/dedmanrecordsandacademicservices/forms/dedman-internship-contract-03262021.pdf"),
        ("University of North Texas - Employer internship information", "https://class.unt.edu/political-science/undergraduate/internships/employer-information.html"),
        ("U.S. Department of Education - FERPA and parents of postsecondary students", "https://studentprivacy.ed.gov/faq/must-postsecondary-institutions-provide-parent-access-eligible-students-education-records"),
        ("U.S. Department of State - Students abroad safety guidance", "https://travel.state.gov/content/travel/en/international-travel/before-you-go/travelers-with-special-considerations/students.htmldisclaimers.html"),
        ("CIEE - 2026 summer global internship fees", "https://www.ciee.org/sites/default/files/documents/2025-06/fee-entry-worksheet-summer-2026.pdf"),
        ("AT3 - 2026 ministry internship fees", "https://at-3.org/faqs-for-interns"),
        ("SEPE International - Three-month Christian missions internship", "https://sepeinternational.org/program/christian-mission-internships/3-months-internship/"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        add_link(p, label, url)

    add_legal_review_note(doc)
    return save(doc, "Hope-Sojourns-Internship-Program-Business-Model.docx")


def build_university_agreement():
    doc = new_document("University Placement Agreement", "Template for a college or university partnership", "Agreement template", compact=True)
    add_legal_review_note(doc)
    add_labeled_fields(doc, [
        ("University", "FULL LEGAL NAME"), ("Hope Sojourns entity", "FULL LEGAL NAME"),
        ("Effective date", "DATE"), ("Initial term", "START AND END DATE"),
        ("University contact", "NAME, TITLE, EMAIL, PHONE"), ("Hope Sojourns contact", "NAME, TITLE, EMAIL, PHONE"),
    ])
    add_heading(doc, "1. Purpose", 1)
    add_paragraph(doc, "This agreement establishes a framework for eligible students to participate in ministry-based experiential learning placements coordinated by Hope Sojourns and supervised by approved host ministries. Individual placements require a separate learning agreement.")
    add_heading(doc, "2. University responsibilities", 1)
    add_bullets(doc, [
        "Determine student eligibility, academic credit, required hours, assignments, faculty oversight, and grades.",
        "Provide Hope Sojourns with applicable placement standards, evaluation forms, deadlines, and accessibility requirements.",
        "Identify a faculty or internship contact for routine communication and incident escalation.",
        "Obtain student consent before sharing protected education records unless another lawful basis applies.",
        "Inform Hope Sojourns promptly of material changes to academic or participation requirements.",
    ])
    add_heading(doc, "3. Hope Sojourns responsibilities", 1)
    add_bullets(doc, [
        "Vet potential host ministries and proposed roles using documented criteria.",
        "Screen applicants and coordinate host interviews without guaranteeing placement.",
        "Provide structured preparation, weekly check-ins, and agreed university updates.",
        "Maintain safeguarding, background-check, emergency-response, privacy, and incident-reporting procedures.",
        "Notify the university promptly of a serious safety, conduct, supervision, or placement concern, subject to law and student privacy.",
        "Coordinate midpoint and final host evaluations when required.",
    ])
    add_heading(doc, "4. Shared placement standards", 1)
    add_bullets(doc, [
        "A written role description, defined dates, expected hours, direct on-site supervisor, and measurable learning objectives",
        "Work that supports learning and does not improperly displace employees",
        "Reasonable access to supervision, necessary resources, and a safe work setting",
        "A written plan for housing, food, local transportation, and arrival when relevant",
        "Compliance with applicable visa, work-permission, insurance, safeguarding, and mandatory-reporting requirements",
    ])
    add_heading(doc, "5. Student selection and removal", 1)
    add_paragraph(doc, "The university may refer or approve academically eligible students. Hope Sojourns and the host ministry retain responsibility for fit assessment and final placement acceptance. Any party may recommend removal when continued participation presents a material safety, conduct, academic, supervision, or mission-fit concern. The parties will coordinate a fair and timely response whenever circumstances allow.")
    add_heading(doc, "6. Financial responsibilities", 1)
    add_paragraph(doc, "The university will disclose its tuition and fees directly to the student. Hope Sojourns will disclose its Program Support Fee and estimated leader travel expenses before the student signs the participation agreement. Unless otherwise documented, the student is responsible for personal travel, housing, food, and local transportation when the host ministry does not provide them.")
    add_heading(doc, "7. Privacy and records", 1)
    add_paragraph(doc, "Each party will limit access to student information, use it only for the placement, protect it using reasonable safeguards, report suspected unauthorized access promptly, and retain or destroy records according to law and the written data schedule. If Hope Sojourns acts as a university service provider, the parties will document any FERPA requirements and permitted uses.")
    add_heading(doc, "8. Safety, incidents, and insurance", 1)
    add_paragraph(doc, "Each party will maintain the insurance required in the final agreement. Hope Sojourns will maintain an emergency plan and identify primary and backup contacts. Serious injury, missing-person concerns, abuse allegations, arrest, evacuation, or other major incidents will be escalated without unreasonable delay.")
    add_heading(doc, "9. Independent parties; no guarantee", 1)
    add_paragraph(doc, "The parties remain independent organizations. This agreement does not create employment, agency, a joint venture, academic credit, or a guarantee of placement. Student employment status and compensation must be evaluated for each placement under applicable law.")
    add_heading(doc, "10. Term, termination, and notices", 1)
    add_paragraph(doc, "Either party may end this agreement upon [NUMBER] days' written notice, or immediately for a material safety, legal, insurance, or policy concern. Ending this agreement does not automatically determine the outcome of an active student placement; the parties will make a student-centered transition plan when reasonably possible.")
    add_heading(doc, "11. Additional terms", 1)
    add_labeled_fields(doc, [
        ("Governing law", "STATE"), ("Notice addresses", "ADDRESS AND EMAIL FOR EACH PARTY"),
        ("Required exhibits", "LEARNING AGREEMENT, INSURANCE, DATA TERMS, OTHER"),
    ])
    add_heading(doc, "Signatures", 1)
    add_signature_table(doc, ["Authorized university representative", "Authorized Hope Sojourns representative"])
    return save(doc, "University-Placement-Agreement-Template.docx")


def build_host_agreement():
    doc = new_document("Host Ministry Agreement", "Template for an approved internship site", "Agreement template", compact=True)
    add_legal_review_note(doc)
    add_labeled_fields(doc, [
        ("Host ministry", "FULL LEGAL NAME"), ("Placement location", "CITY, STATE/COUNTRY"),
        ("On-site supervisor", "NAME, TITLE, EMAIL, PHONE"), ("Hope Sojourns coordinator", "NAME, EMAIL, PHONE"),
        ("Agreement term", "START AND END DATE"),
    ])
    add_heading(doc, "1. Purpose and placement", 1)
    add_paragraph(doc, "The host ministry agrees to provide a supervised, educational internship experience for students accepted through Hope Sojourns. Each student will have a separate role description and learning agreement.")
    add_heading(doc, "2. Host ministry commitments", 1)
    add_bullets(doc, [
        "Name a qualified supervisor who is available for orientation, regular meetings, feedback, and evaluations.",
        "Provide meaningful duties connected to the learning agreement and avoid using the intern merely to replace paid staff.",
        "Provide a safe work setting, necessary tools, local orientation, and clear schedules and boundaries.",
        "Explain workplace, confidentiality, media, technology, safeguarding, and incident-reporting rules.",
        "Report serious safety, conduct, performance, supervision, housing, or legal concerns to Hope Sojourns promptly.",
        "Cooperate with reasonable university and Hope Sojourns evaluation requirements.",
    ])
    add_heading(doc, "3. Support provided by the host", 1)
    add_checkbox_list(doc, [
        "Housing provided at no cost",
        "Housing subsidized; student cost: [AMOUNT AND PAYMENT TERMS]",
        "Meals provided: [DETAILS]",
        "Food allowance or stipend provided: [DETAILS]",
        "Local transportation provided: [DETAILS]",
        "Intern stipend or wage provided: [DETAILS]",
        "No housing, food, local transportation, or stipend provided",
    ])
    add_heading(doc, "4. Safeguarding and vulnerable people", 1)
    add_paragraph(doc, "Before an intern works with children, trafficking survivors, refugees, people experiencing homelessness, or other vulnerable people, the host will identify required screening, training, supervision, consent, confidentiality, and mandatory-reporting procedures. The stricter applicable Hope Sojourns, university, host, or local requirement will control unless counsel directs otherwise.")
    add_heading(doc, "5. International requirements", 1)
    add_bullets(doc, [
        "Confirm the lawful visa, volunteer, internship, or work-permission pathway before travel.",
        "Identify a local emergency contact, nearby medical resources, and safe transportation options.",
        "Explain relevant local laws, customs, security considerations, and limits on ministry activity.",
        "Provide accurate housing and arrival information before the student makes nonrefundable purchases.",
    ])
    add_heading(doc, "6. Communication and evaluation", 1)
    add_labeled_fields(doc, [
        ("Supervisor meeting frequency", "AT LEAST WEEKLY OR OTHER"),
        ("Midpoint evaluation due", "DATE"), ("Final evaluation due", "DATE"),
        ("Routine host/Hope Sojourns check-in", "FREQUENCY AND METHOD"),
    ])
    add_heading(doc, "7. Insurance and responsibility", 1)
    add_paragraph(doc, "The parties will confirm applicable general liability, professional liability, abuse/molestation, workers' compensation, automobile, travel, and other coverage before placement. Each party remains responsible for its own acts and omissions to the extent provided by law and the final reviewed agreement.")
    add_heading(doc, "8. Ending or pausing a placement", 1)
    add_paragraph(doc, "The host may pause duties immediately when necessary to protect safety or comply with law. Except in urgent situations, the host will consult Hope Sojourns and the university before ending a placement. The parties will document reasons, student support, travel or housing implications, and required notifications.")
    add_heading(doc, "9. Signatures", 1)
    add_signature_table(doc, ["Authorized host ministry representative", "Authorized Hope Sojourns representative"])
    return save(doc, "Host-Ministry-Agreement-Template.docx")


def build_student_agreement():
    doc = new_document("Student Participation and Financial Agreement", "Template for an accepted internship participant", "Participant agreement", compact=True)
    add_legal_review_note(doc)
    add_labeled_fields(doc, [
        ("Student", "FULL LEGAL NAME"), ("University and program", "SCHOOL, DEPARTMENT, DEGREE"),
        ("Host ministry", "NAME"), ("Placement", "CITY, STATE/COUNTRY"),
        ("Placement dates", "START AND END DATE"), ("Hope Sojourns coordinator", "NAME AND CONTACT"),
    ])
    add_heading(doc, "1. Participation", 1)
    add_paragraph(doc, "I understand that acceptance is contingent on final host approval, required agreements, academic approval when applicable, payment, background screening, insurance, travel documents, and continued compliance with program policies.")
    add_heading(doc, "2. Program support", 1)
    add_bullets(doc, [
        "Placement coordination and university learning-agreement support",
        "Three preparation sessions",
        "Weekly Hope Sojourns check-ins during the placement",
        "Routine communication with the host and university",
        "Arrival planning and leader accompaniment when specifically included",
        "Incident escalation and transition support within the program's stated capacity",
    ])
    add_heading(doc, "3. Financial summary", 1)
    add_labeled_fields(doc, [
        ("Program Support Fee", "$AMOUNT"), ("Application/matching deposit credited to fee", "$AMOUNT"),
        ("Estimated Hope Sojourns leader travel share", "$AMOUNT; MAXIMUM AUTHORIZED $AMOUNT"),
        ("Payment schedule", "DATES AND AMOUNTS"), ("Scholarship or subsidy", "$AMOUNT AND SOURCE"),
    ])
    add_heading(doc, "4. Student-paid expenses", 1)
    add_bullets(doc, [
        "University tuition and fees, if required",
        "Personal airfare and other travel",
        "Housing, food, and local transportation unless the host ministry provides or subsidizes them in writing",
        "Passport, visa, screening, vaccination, insurance, personal, and emergency expenses",
        "The Program Support Fee and disclosed leader travel share",
    ])
    add_paragraph(doc, "Third-party travel, housing, visa, insurance, or vendor charges are subject to the provider's terms and may be nonrefundable. I will not make nonrefundable purchases until Hope Sojourns confirms that I may proceed.")
    add_heading(doc, "5. Student commitments", 1)
    add_bullets(doc, [
        "Complete university assignments and hours independently; Hope Sojourns does not award credit or grades.",
        "Attend preparation and weekly check-ins and respond promptly to reasonable communications.",
        "Follow the Code of Conduct, Safeguarding Policy, host rules, and applicable law.",
        "Respect confidentiality, obtain consent before photos or stories, and protect the dignity of people served.",
        "Disclose material health, accessibility, legal, or safety information through the approved confidential process when necessary for planning.",
        "Report incidents, abuse concerns, harassment, threats, arrest, missing-person concerns, or significant illness promptly.",
    ])
    add_heading(doc, "6. Travel and insurance", 1)
    add_checkbox_list(doc, [
        "I have a valid passport and required visa/work authorization or will obtain them by [DATE].",
        "I will maintain required travel medical, medical evacuation, repatriation, and trip-protection coverage.",
        "I will provide policy and emergency-assistance information before departure.",
        "I understand that the U.S. government generally does not pay my overseas medical expenses or evacuation costs.",
    ])
    add_heading(doc, "7. Risk and emergency response", 1)
    add_paragraph(doc, "Travel and ministry service involve risks that cannot be eliminated. The final attorney-reviewed agreement should include appropriate risk acknowledgments, releases, governing law, dispute terms, medical-consent language, and emergency authority. I agree to follow the trip-specific emergency plan and reasonable safety instructions.")
    add_heading(doc, "8. Cancellation, removal, and refunds", 1)
    add_paragraph(doc, "The current Cancellation and Refund Policy is incorporated by reference. Academic failure, voluntary withdrawal, misconduct, inaccurate application information, visa denial, travel disruption, host closure, or safety conditions may affect the placement and available refund. Hope Sojourns will document the applicable outcome.")
    add_heading(doc, "9. Acknowledgment and signatures", 1)
    add_checkbox_list(doc, [
        "I received the final cost summary and understand what the local ministry will and will not provide.",
        "I received the Code of Conduct, safeguarding information, refund policy, privacy notice, and emergency plan.",
        "I had an opportunity to ask questions and will seek independent advice if needed.",
    ])
    add_signature_table(doc, ["Student participant", "Hope Sojourns representative", "Parent/guardian only if legally required"])
    return save(doc, "Student-Participation-and-Financial-Agreement-Template.docx")


def build_learning_agreement():
    doc = new_document("Internship Learning Agreement", "Shared plan for the student, university, Hope Sojourns, and host", "Learning template", compact=True)
    add_labeled_fields(doc, [
        ("Student", "NAME, EMAIL, PHONE"), ("University", "SCHOOL, DEPARTMENT, COURSE"),
        ("Faculty sponsor", "NAME, EMAIL, PHONE"), ("Host ministry", "NAME AND LOCATION"),
        ("Site supervisor", "NAME, TITLE, EMAIL, PHONE"), ("Hope Sojourns mentor", "NAME, EMAIL, PHONE"),
        ("Dates", "START AND END"), ("Expected schedule", "DAYS, HOURS/WEEK, TOTAL HOURS"),
        ("Position title", "TITLE"),
    ])
    add_heading(doc, "Placement purpose", 1)
    add_paragraph(doc, "[Describe how the role connects the student's academic program, vocational interests, ministry context, and the host's genuine needs.]")
    add_heading(doc, "Learning objectives", 1)
    add_table(doc, ["Objective", "Activities and evidence", "How progress will be evaluated"], [
        ("1. [Knowledge or skill]", "[Tasks, observation, training, deliverable]", "[Supervisor feedback, portfolio, reflection, rubric]"),
        ("2. [Knowledge or skill]", "[Tasks, observation, training, deliverable]", "[Supervisor feedback, portfolio, reflection, rubric]"),
        ("3. [Knowledge or skill]", "[Tasks, observation, training, deliverable]", "[Supervisor feedback, portfolio, reflection, rubric]"),
        ("4. [Knowledge or skill]", "[Tasks, observation, training, deliverable]", "[Supervisor feedback, portfolio, reflection, rubric]"),
    ], [2300, 4060, 3000], font_size=8.7)
    add_heading(doc, "Core responsibilities", 1)
    add_bullets(doc, [
        "[Responsibility connected to learning objective]",
        "[Responsibility connected to learning objective]",
        "[Project or final deliverable]",
        "[Reflection, reporting, or presentation requirement]",
    ])
    add_heading(doc, "Supervision and communication", 1)
    add_table(doc, ["Touchpoint", "Frequency", "Owner", "Documentation"], [
        ("On-site supervision", "[Daily/weekly]", "Host supervisor", "[Notes or task system]"),
        ("Hope Sojourns mentoring", "Weekly", "Hope Sojourns mentor", "Weekly check-in record"),
        ("Faculty contact", "[Frequency]", "Faculty sponsor", "[Course platform/email]"),
        ("Midpoint evaluation", "[Date]", "Host + student", "University form"),
        ("Final evaluation", "[Date]", "Host + student", "University form"),
    ], [2600, 1700, 2300, 2760], font_size=8.8)
    add_heading(doc, "Boundaries and special requirements", 1)
    add_labeled_fields(doc, [
        ("Confidentiality requirements", "DETAILS"), ("Safeguarding/screening", "DETAILS"),
        ("Physical or travel requirements", "DETAILS"), ("Approved accommodations", "DETAILS OR NONE"),
        ("Prohibited or out-of-scope duties", "DETAILS"),
    ])
    # Keep the four-party approval block together on a clean signature page.
    doc.add_page_break()
    add_heading(doc, "Agreement", 1)
    add_paragraph(doc, "The signers understand their roles, will communicate material changes promptly, and will use this plan as the basis for supervision and evaluation. Academic credit and grades remain under university control.")
    add_signature_table(doc, ["Student", "Host site supervisor", "Faculty sponsor", "Hope Sojourns representative"])
    return save(doc, "Internship-Learning-Agreement-Template.docx")


def build_code_of_conduct():
    doc = new_document("Internship Code of Conduct", "Behavioral expectations for every Hope Sojourns participant", "Program policy", status="Draft for adoption", compact=True)
    add_legal_review_note(doc)
    add_heading(doc, "Purpose", 1)
    add_paragraph(doc, "Hope Sojourns internships place students in communities where trust, dignity, safety, cultural humility, and faithful service matter. Participants are expected to represent Hope Sojourns, their university, and the host ministry responsibly.")
    sections = [
        ("Respect and dignity", [
            "Treat every person with dignity regardless of nationality, race, ethnicity, sex, disability, age, economic status, religion, immigration history, or life experience.",
            "Avoid language or behavior that humiliates, stereotypes, exploits, pressures, or objectifies another person.",
            "Listen before acting and follow local leaders when deciding what help is appropriate.",
        ]),
        ("Safeguarding and boundaries", [
            "Follow all child, vulnerable-adult, survivor-care, and mandatory-reporting requirements.",
            "Do not pursue romantic or sexual relationships with people receiving services, minors, or anyone over whom the participant has power or influence.",
            "Do not meet, transport, photograph, message, or remain alone with a minor or vulnerable person outside approved procedures.",
        ]),
        ("Harassment, violence, and retaliation", [
            "Harassment, discrimination, bullying, sexual misconduct, threats, violence, hazing, and retaliation are prohibited.",
            "Report concerns promptly using the safeguarding and incident process. Immediate danger should be reported to local emergency services first.",
        ]),
        ("Alcohol, drugs, and illegal activity", [
            "Illegal drugs, misuse of medication, intoxication during program duties, and conduct that violates local law or host policy are prohibited.",
            "When alcohol is lawful, participants must follow the stricter Hope Sojourns, university, host, or local standard and must never allow alcohol use to impair judgment or safety.",
        ]),
        ("Confidentiality, stories, and media", [
            "Protect personal, medical, legal, financial, ministry, and university information.",
            "Obtain informed permission before taking or sharing identifiable photos, recordings, or stories.",
            "Never reveal information that could endanger a survivor, refugee, person experiencing homelessness, recovering individual, child, or ministry partner.",
        ]),
        ("Attendance and professionalism", [
            "Be punctual, prepared, sober, appropriately dressed, and honest about availability and progress.",
            "Use ministry property, vehicles, funds, technology, and supplies only as authorized.",
            "Inform the supervisor and Hope Sojourns promptly about illness, absence, conflict, or inability to perform an assigned duty.",
        ]),
        ("Travel and personal safety", [
            "Follow approved transportation, lodging, communication, and check-in procedures.",
            "Do not enter prohibited areas, ride with an unauthorized driver, or change lodging or travel plans without required notice.",
            "Comply with local law, visa conditions, travel advisories, and reasonable safety instructions.",
        ]),
    ]
    for title, items in sections:
        add_heading(doc, title, 2)
        add_bullets(doc, items)
    add_heading(doc, "Concern and response", 1)
    add_paragraph(doc, "Concerns may lead to coaching, a corrective plan, temporary removal from duties, relocation, dismissal, notification of the university, or emergency action depending on severity. Hope Sojourns will seek a fair response while prioritizing immediate safety and legal obligations.")
    add_heading(doc, "Participant acknowledgment", 1)
    add_checkbox_list(doc, [
        "I have read and understand this Code of Conduct.",
        "I know how to report a safety, safeguarding, harassment, or misconduct concern.",
        "I understand that serious or repeated violations may end my placement without refund beyond the published policy.",
    ])
    add_signature_table(doc, ["Student participant", "Hope Sojourns representative"])
    return save(doc, "Internship-Code-of-Conduct.docx")


def build_safeguarding_policy():
    doc = new_document("Safeguarding and Misconduct Reporting Policy", "Protection, reporting, response, and non-retaliation standards", "Program policy", status="Draft for adoption", compact=True)
    add_legal_review_note(doc)
    add_heading(doc, "1. Policy commitment", 1)
    add_paragraph(doc, "Hope Sojourns is committed to protecting children, vulnerable adults, survivors, students, staff, volunteers, ministry partners, and community members from abuse, exploitation, harassment, neglect, retaliation, and avoidable harm.")
    add_heading(doc, "2. Scope", 1)
    add_paragraph(doc, "This policy applies to Hope Sojourns leaders, employees, volunteers, interns, contractors, and participants during selection, preparation, travel, service, housing, transportation, communications, and program-related online activity.")
    add_heading(doc, "3. Prohibited conduct", 1)
    add_bullets(doc, [
        "Physical, sexual, emotional, financial, spiritual, or online abuse or exploitation",
        "Grooming, sexualized communication, coercion, trafficking, commercial sexual activity, or exchange of help for attention or favors",
        "Harassment, discrimination, bullying, threats, violence, hazing, stalking, or retaliation",
        "Unapproved one-on-one isolation, transportation, lodging, or private digital communication with a minor or vulnerable person",
        "Sharing confidential information, images, or stories without lawful and informed authorization",
        "Ignoring, concealing, discouraging, or privately investigating a report that should be escalated",
    ])
    add_heading(doc, "4. Prevention standards", 1)
    add_bullets(doc, [
        "Role-specific screening and background checks before access to minors or vulnerable people",
        "Safeguarding orientation and written acknowledgment before service begins",
        "At least two authorized adults present or another approved observable-accountable practice when working with minors",
        "Clear rules for transportation, bathrooms, changing areas, lodging, photography, gifts, money, and electronic communication",
        "A named local safeguarding lead and a Hope Sojourns primary and backup reporting contact",
        "Periodic review of housing, work sites, supervision, and foreseeable risks",
    ])
    add_heading(doc, "5. How to report", 1)
    add_callout(doc, "Immediate danger", "Call local emergency services first. Then contact the local safeguarding lead and Hope Sojourns emergency contacts as soon as it is safe to do so.", fill="FFF0EC", accent=RED)
    add_labeled_fields(doc, [
        ("Hope Sojourns primary safeguarding contact", "NAME, PHONE, EMAIL"),
        ("Hope Sojourns backup safeguarding contact", "NAME, PHONE, EMAIL"),
        ("Host safeguarding contact", "NAME, PHONE, EMAIL"),
        ("University contact", "NAME, PHONE, EMAIL"),
        ("Local emergency/mandatory-reporting contacts", "NUMBERS AND AGENCIES"),
    ])
    add_paragraph(doc, "A report should include only known facts: who was involved, what was observed or disclosed, when and where it occurred, immediate safety needs, witnesses, and actions already taken. Do not promise secrecy; explain that information will be shared only with people who need it for protection, response, or legal compliance.")
    add_heading(doc, "6. Response protocol", 1)
    add_numbered(doc, [
        "Protect: address immediate medical, physical, and emotional safety without confronting an alleged offender unnecessarily.",
        "Preserve: document exact words and observable facts; preserve messages, images, and other evidence lawfully.",
        "Report: notify the designated contacts and authorities required by applicable law and policy.",
        "Separate: restrict duties, contact, transportation, housing, or system access when needed to prevent further harm.",
        "Support: offer appropriate medical, advocacy, counseling, university, pastoral, or travel support without pressuring the affected person.",
        "Coordinate: determine university, insurer, board, legal, family, and government notifications with qualified advisors.",
        "Review: document findings, actions, lessons learned, and required policy or partner changes.",
    ])
    add_heading(doc, "7. Non-retaliation and confidentiality", 1)
    add_paragraph(doc, "Good-faith reports and participation in a response process are protected from retaliation. Deliberately false reports may be addressed separately, but an unsubstantiated report is not automatically a false report. Information will be limited to those with a legitimate need to know.")
    add_heading(doc, "8. Initial incident record", 1)
    add_labeled_fields(doc, [
        ("Reporter and safe contact method", "DETAILS"), ("Date/time received", "DATE AND TIME"),
        ("People involved", "NAMES/IDENTIFIERS"), ("Location", "DETAILS"),
        ("Observed or disclosed facts", "USE ATTACHED FACTUAL NOTES"), ("Immediate safety action", "DETAILS"),
        ("Authorities/contacts notified", "WHO, WHEN, BY WHOM"), ("Next review time", "DATE AND TIME"),
    ])
    return save(doc, "Safeguarding-and-Misconduct-Reporting-Policy.docx")


def build_background_policy():
    doc = new_document("Background Check Policy", "Screening standards for internship participants and leaders", "Program policy", status="Draft for adoption", compact=True)
    add_legal_review_note(doc)
    add_heading(doc, "1. Purpose", 1)
    add_paragraph(doc, "Background screening supports informed placement decisions and helps protect students, ministries, children, vulnerable adults, and communities. A background check is one part of screening and does not replace references, interviews, training, supervision, or safeguarding controls.")
    add_heading(doc, "2. Who is screened", 1)
    add_bullets(doc, [
        "Hope Sojourns leaders who travel with or supervise students",
        "Interns whose duties involve children, vulnerable adults, survivor services, homes, finances, transportation, or unsupervised access",
        "Other staff, volunteers, or contractors when required by law, a university, insurer, or host ministry",
    ])
    add_heading(doc, "3. Screening components", 1)
    add_checkbox_list(doc, [
        "Identity and address-history verification",
        "National and relevant county/state criminal history",
        "National sex-offender registry search",
        "Motor vehicle record when driving is part of the role",
        "Professional license or credential verification when relevant",
        "International or host-country screening where lawful and reasonably available",
        "Two references, including one who has observed the applicant in a service, work, school, or ministry setting",
    ])
    add_heading(doc, "4. Consent and fair process", 1)
    add_paragraph(doc, "Hope Sojourns will obtain required written consent and provide disclosures separately when a consumer-reporting agency is used. Any preliminary adverse decision will follow applicable notice, copy, dispute, and waiting-period requirements. Screening will be applied consistently and reviewed individually rather than through an automatic blanket exclusion unless law requires otherwise.")
    add_heading(doc, "5. Review factors", 1)
    add_bullets(doc, [
        "Nature and seriousness of the conduct",
        "Time elapsed and evidence of rehabilitation",
        "Accuracy and completeness of the applicant's disclosure",
        "Relationship between the record and the proposed duties, access, population, travel, housing, or driving",
        "Applicable university, insurer, host, legal, and safeguarding requirements",
        "Whether supervision or role modification can reasonably reduce risk",
    ])
    add_heading(doc, "6. Timing and renewal", 1)
    add_paragraph(doc, "Screening must be complete before access begins. Re-screen active leaders and recurring participants every [TWO/THREE] years, and sooner after a material role change, credible report, legal requirement, or break in service of [NUMBER] months.")
    add_heading(doc, "7. Records and confidentiality", 1)
    add_paragraph(doc, "Store authorization, reports, adjudication notes, and results separately from routine program records with access limited to designated decision-makers. Share only the minimum necessary result, such as eligible, eligible with conditions, or not eligible. Destroy records securely according to the adopted retention schedule and applicable law.")
    # Start the fill-in decision record as one labeled form block.
    doc.add_page_break()
    add_heading(doc, "8. Screening decision record", 1)
    add_labeled_fields(doc, [
        ("Applicant", "NAME"), ("Role/placement", "DETAILS"), ("Screening date/provider", "DETAILS"),
        ("Decision", "ELIGIBLE / ELIGIBLE WITH CONDITIONS / NOT ELIGIBLE"),
        ("Conditions or role limits", "DETAILS"), ("Reviewers", "NAMES AND TITLES"),
        ("Next screening due", "DATE"),
    ])
    return save(doc, "Background-Check-Policy.docx")


def build_refund_policy():
    doc = new_document("Cancellation and Refund Policy", "Draft participant-facing terms for internship placements", "Program policy", status="Draft values for review", compact=True)
    add_legal_review_note(doc)
    add_callout(doc, "Decision required", "The percentages and deadlines below are a reasonable pilot proposal, not an adopted promise. Confirm them against actual staffing, vendor, travel, insurance, and accounting practices before publication.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "1. Application and matching deposit", 1)
    add_paragraph(doc, "A $250 application and matching deposit is due when Hope Sojourns begins individualized placement work. The deposit is credited to the Program Support Fee. It is refundable if Hope Sojourns cannot offer an approved placement, but otherwise becomes nonrefundable after the first host introduction or [NUMBER] hours of individualized work, whichever occurs first.")
    add_heading(doc, "2. Program Support Fee refunds", 1)
    add_table(doc, ["Student cancellation received", "Proposed refund of Program Support Fee paid"], [
        ("Before an approved placement is confirmed", "100%, less any earned/nonrefundable deposit disclosed above"),
        ("60 or more days before departure/start", "75%"),
        ("30-59 days before departure/start", "50%"),
        ("15-29 days before departure/start", "25%"),
        ("14 days or fewer before departure/start", "No refund"),
        ("After travel or placement begins", "No refund, except a documented discretionary amount when Hope Sojourns or the host materially fails to provide the agreed program"),
    ], [4800, 4560], font_size=9)
    add_heading(doc, "3. Third-party and travel costs", 1)
    add_paragraph(doc, "Airfare, lodging, insurance, visas, background checks, local transportation, deposits, currency losses, and other third-party costs follow the vendor's terms. Nonrefundable or already-spent Hope Sojourns leader travel costs remain the student's responsibility up to the written authorized maximum unless Hope Sojourns caused the cancellation through material breach.")
    add_heading(doc, "4. Host, university, visa, or safety cancellation", 1)
    add_bullets(doc, [
        "If Hope Sojourns cannot provide the confirmed placement before departure, it will offer a reasonable alternative or refund the unused Program Support Fee.",
        "University denial of credit does not automatically cancel the nonacademic placement unless academic credit was expressly made a condition in writing.",
        "Visa denial, passport delay, illness, family emergency, or travel disruption is treated as a student cancellation unless a written exception applies; participants should consider trip-cancellation coverage.",
        "Hope Sojourns may change, pause, relocate, or cancel a placement for safety, legal, insurance, disaster, host, or government reasons. The unused Program Support Fee will be evaluated in good faith after nonrecoverable obligations are identified.",
    ])
    add_heading(doc, "5. Dismissal", 1)
    add_paragraph(doc, "A participant dismissed for serious misconduct, material application misrepresentation, unlawful activity, refusal to follow safety instructions, or repeated policy violations is not entitled to a Program Support Fee refund and remains responsible for safe return travel and other resulting costs, subject to applicable law and the final agreement.")
    add_heading(doc, "6. Request process", 1)
    add_paragraph(doc, "Cancellation and refund requests must be submitted in writing to [EMAIL] and include the participant's name, placement, reason, requested effective date, and supporting documentation when relevant. Hope Sojourns will issue a written determination within [10] business days and approved refunds within [30] days.")
    add_heading(doc, "Participant acknowledgment", 1)
    add_signature_table(doc, ["Student participant", "Hope Sojourns representative"])
    return save(doc, "Cancellation-and-Refund-Policy.docx")


def build_emergency_plan():
    doc = new_document("Emergency Response Plan", "Master framework and trip-specific planning template", "Safety template", status="Draft for adoption", compact=True)
    add_legal_review_note(doc)
    add_heading(doc, "1. Core contacts", 1)
    add_labeled_fields(doc, [
        ("Hope Sojourns incident lead", "NAME, 24/7 PHONE, EMAIL"),
        ("Hope Sojourns backup lead", "NAME, 24/7 PHONE, EMAIL"),
        ("Host emergency contact", "NAME, 24/7 PHONE"),
        ("University emergency contact", "NAME, PHONE, EMAIL"),
        ("Student emergency contact", "NAME, RELATIONSHIP, PHONE"),
        ("Insurance assistance", "CARRIER, POLICY, 24/7 PHONE"),
        ("Local emergency services", "POLICE/FIRE/MEDICAL NUMBERS"),
        ("Nearest hospital/clinic", "NAME, ADDRESS, PHONE"),
        ("U.S. embassy/consulate", "LOCATION, PHONE"),
    ])
    add_heading(doc, "2. Response priorities", 1)
    add_numbered(doc, [
        "Protect life and address immediate danger.",
        "Contact local emergency services and obtain medical care when needed.",
        "Account for the student and prevent additional harm.",
        "Notify the Hope Sojourns primary or backup incident lead.",
        "Preserve factual information and evidence without conducting an amateur investigation.",
        "Coordinate host, university, insurer, family, legal, and government notifications.",
        "Plan safe housing, transportation, evacuation, return, or continued participation.",
        "Document actions, follow-up, and lessons learned.",
    ])
    add_heading(doc, "3. Incident levels", 1)
    add_table(doc, ["Level", "Examples", "Initial response"], [
        ("1 - Routine", "Minor illness, missed check-in resolved quickly, ordinary conflict", "Host/Hope Sojourns manage and document during normal operations."),
        ("2 - Significant", "Medical visit, credible harassment concern, housing failure, arrest risk, sustained loss of contact", "Notify incident lead promptly; make a written response plan and university notification decision."),
        ("3 - Critical", "Life-threatening injury, sexual assault, abuse allegation, missing student, violent crime, evacuation, death", "Call emergency services; activate primary and backup leadership; notify insurer, university, family, counsel, and authorities as required."),
    ], [1350, 4310, 3700], font_size=8.6)
    add_heading(doc, "4. Scenario checklists", 1)
    scenarios = {
        "Medical or mental-health emergency": ["Obtain emergency care", "Contact insurer/assistance provider", "Secure medications and safe supervision", "Coordinate family/university communication", "Document fitness-to-return decision"],
        "Missing or unreachable participant": ["Confirm last known contact and itinerary", "Check approved contacts/locations", "Contact local authorities when threshold is met", "Preserve phone and travel information lawfully", "Escalate to family, university, and embassy as appropriate"],
        "Crime, assault, abuse, or harassment": ["Move to safety and obtain medical/advocacy support", "Preserve choices and evidence", "Follow safeguarding and mandatory-reporting rules", "Prevent contact or retaliation", "Coordinate trauma-informed follow-up"],
        "Civil unrest, disaster, or evacuation": ["Monitor authoritative guidance", "Account for all participants", "Shelter or relocate", "Contact insurer and embassy/STEP resources", "Approve evacuation or return travel"],
        "Housing or transportation failure": ["Move the student to a verified safe temporary option", "Prevent use of an unsafe driver or location", "Notify host and Hope Sojourns", "Document replacement cost and responsibilities", "Update university/family if material"],
    }
    for title, items in scenarios.items():
        add_heading(doc, title, 2)
        add_checkbox_list(doc, items)
    add_heading(doc, "5. Communication rules", 1)
    add_bullets(doc, [
        "Use one designated incident lead to reduce conflicting information.",
        "Share verified facts, actions, next update time, and a safe contact method.",
        "Do not identify affected people publicly or speak to media without authorization.",
        "Follow the student's parent/emergency-contact authorization except when immediate safety, law, or emergency conditions require another response.",
        "Record who was notified, by whom, when, what was shared, and what follow-up was promised.",
    ])
    # Keep the readiness checklist and incident log together as a field-use section.
    doc.add_page_break()
    add_heading(doc, "6. Trip-specific readiness check", 1)
    add_checkbox_list(doc, [
        "Travel advisory and country information reviewed",
        "Passport/visa/work authorization confirmed",
        "STEP enrollment completed for international placement",
        "Travel medical, evacuation, repatriation, and assistance coverage verified",
        "Arrival, housing, local transportation, and backup lodging verified",
        "Participant medical/accessibility plan reviewed confidentially",
        "Primary and backup communication methods tested",
        "Emergency funds/payment method available",
        "Host and university received current contact sheet",
    ])
    add_heading(doc, "7. Incident log", 1)
    add_table(doc, ["Date/time", "Known facts", "Actions and notifications", "Owner/next update"], [
        ("[DATE/TIME]", "[FACTUAL SUMMARY]", "[ACTION, PERSON CONTACTED, TIME]", "[NAME/TIME]"),
        ("[DATE/TIME]", "[FACTUAL SUMMARY]", "[ACTION, PERSON CONTACTED, TIME]", "[NAME/TIME]"),
        ("[DATE/TIME]", "[FACTUAL SUMMARY]", "[ACTION, PERSON CONTACTED, TIME]", "[NAME/TIME]"),
    ], [1500, 2900, 3300, 1660], font_size=8.5)
    return save(doc, "Emergency-Response-Plan-Template.docx")


def build_parent_authorization():
    doc = new_document("Parent and Emergency Contact Authorization", "Student-directed communication preferences and emergency consent", "Authorization template", compact=True)
    add_legal_review_note(doc)
    add_labeled_fields(doc, [
        ("Student", "FULL LEGAL NAME"), ("Placement", "HOST AND LOCATION"),
        ("Placement dates", "DATES"), ("University", "NAME"),
        ("Primary emergency contact", "NAME, RELATIONSHIP, PHONE, EMAIL"),
        ("Secondary emergency contact", "NAME, RELATIONSHIP, PHONE, EMAIL"),
    ])
    add_heading(doc, "Student direction", 1)
    add_paragraph(doc, "I understand that I control routine disclosure of my personal and program information as an adult student, subject to applicable law, university rules, contractual obligations, and emergency exceptions. I direct Hope Sojourns as follows:")
    add_checkbox_list(doc, [
        "Hope Sojourns may confirm my safe arrival and general participation status to the contacts named above.",
        "Hope Sojourns may discuss routine travel, housing, or logistical concerns with the contacts named above.",
        "Hope Sojourns may discuss general well-being concerns that do not include protected academic, medical, counseling, disciplinary, or safeguarding details.",
        "Hope Sojourns may discuss a significant academic or performance concern only after first attempting to speak with me, unless immediate action is needed.",
        "Hope Sojourns may not provide routine updates; contact is authorized only for a serious or emergency concern.",
        "Additional direction: [DETAILS]",
    ])
    add_heading(doc, "Emergency authorization", 1)
    add_paragraph(doc, "I authorize Hope Sojourns to contact the people listed above when it reasonably believes I face a serious health, safety, missing-person, legal, evacuation, housing, or travel emergency, or when I am unable to communicate. Hope Sojourns may share the minimum information reasonably necessary to obtain help, coordinate care or travel, or protect me or others.")
    add_heading(doc, "Limits", 1)
    add_bullets(doc, [
        "This form does not require a university, medical provider, insurer, law-enforcement agency, or host ministry to release information beyond what law and its policies permit.",
        "Hope Sojourns will not promise parents or contacts access to academic records, counseling details, medical records, safeguarding reports, or disciplinary information.",
        "I may revise or revoke routine authorization in writing, but revocation does not affect prior disclosures or lawful emergency actions.",
        "This authorization expires [30] days after the placement ends unless revoked earlier or another date is written here: [DATE].",
    ])
    add_heading(doc, "Signatures", 1)
    add_signature_table(doc, ["Student participant", "Hope Sojourns representative", "Primary emergency contact acknowledgment (optional)"])
    return save(doc, "Parent-and-Emergency-Contact-Authorization.docx")


def build_privacy_policy():
    doc = new_document("Data Privacy and Record Retention Policy", "Collection, access, sharing, security, and disposal standards", "Program policy", status="Draft for adoption", compact=True)
    add_legal_review_note(doc)
    add_heading(doc, "1. Purpose and scope", 1)
    add_paragraph(doc, "This policy governs internship applicant, participant, emergency-contact, university, host ministry, screening, payment, incident, and evaluation information maintained by Hope Sojourns in paper, email, cloud, portal, device, and third-party systems.")
    add_heading(doc, "2. Information collected", 1)
    add_bullets(doc, [
        "Identity, contact, education, skills, availability, interests, references, and application responses",
        "Placement, learning objectives, evaluations, communications, attendance, and check-in notes",
        "Payment status and transaction references, but not unnecessary payment-card data",
        "Travel documents, itinerary, insurance, emergency contacts, and limited health/accessibility information needed for safety planning",
        "Background-check authorization, status, and restricted adjudication records",
        "Incident, safeguarding, conduct, legal, or insurance records",
        "Website and portal security logs needed to protect systems and investigate misuse",
    ])
    add_heading(doc, "3. Permitted uses", 1)
    add_bullets(doc, [
        "Assess and arrange placements",
        "Coordinate academic and host requirements",
        "Prepare, mentor, supervise, evaluate, and support participants",
        "Process fees and document scholarships or subsidies",
        "Protect health, safety, systems, people, property, and legal rights",
        "Meet insurance, accounting, tax, university, contractual, and legal obligations",
        "Improve the program using de-identified or appropriately limited information",
    ])
    add_heading(doc, "4. Access and sharing", 1)
    add_paragraph(doc, "Access is limited to people with a documented need. Hope Sojourns will share the minimum necessary information with the student, university, host, insurer, emergency contact, service provider, counsel, or authority based on consent, contract, legitimate program need, emergency, or legal obligation. Education records received while acting for a university will be handled according to the university agreement and applicable FERPA requirements.")
    add_heading(doc, "5. Security standards", 1)
    add_bullets(doc, [
        "Unique accounts, strong passwords, multi-factor authentication where available, and prompt access removal",
        "Encryption in transit and at rest where supported",
        "Restricted storage for health, screening, safeguarding, legal, and identity documents",
        "Approved devices, current software, secure backups, and periodic access review",
        "No sensitive participant data in unapproved personal messaging, shared drives, or public links",
        "Prompt reporting and documented response to suspected loss, unauthorized access, or disclosure",
    ])
    add_heading(doc, "6. Proposed retention schedule", 1)
    add_table(doc, ["Record category", "Proposed retention", "Notes"], [
        ("Unsuccessful or incomplete applications", "2 years", "Retain only what is useful for follow-up, defense, or reporting."),
        ("Participant agreements, learning records, evaluations", "7 years after completion", "Longer when university contract or law requires."),
        ("Financial and tax records", "7 years", "Coordinate with accountant and payment provider."),
        ("Background-check reports", "Shortest lawful operational period; target 5 years or less", "Store separately; retain eligibility outcome longer only if justified."),
        ("Routine check-in notes", "3 years after completion", "Move material incidents to the incident file."),
        ("Incident, safeguarding, claim, or litigation records", "At least 7 years after closure; longer for minors/claims", "Set final period with counsel and insurer."),
        ("Emergency travel documents", "90 days after safe return unless incident or claim requires longer", "Delete passport/visa copies as soon as operationally unnecessary."),
        ("Portal security logs", "1 year", "Longer when used for an investigation."),
    ], [2950, 2100, 4310], font_size=8.4)
    add_heading(doc, "7. Individual requests and correction", 1)
    add_paragraph(doc, "Hope Sojourns will provide a practical process to request access, correction, restriction, or deletion where applicable. Requests will be verified, documented, and answered within [30] days unless law or a university agreement requires another timeline. Some records cannot be deleted while required for safety, contract, tax, claims, or legal obligations.")
    add_heading(doc, "8. Breach and incident response", 1)
    add_numbered(doc, [
        "Contain access and preserve evidence.",
        "Notify the privacy and incident leads.",
        "Determine what data, people, systems, and jurisdictions are affected.",
        "Consult counsel, insurer, university, vendors, and law enforcement as appropriate.",
        "Provide required notices and practical protection steps without unreasonable delay.",
        "Document root cause, remediation, and prevention measures.",
    ])
    add_heading(doc, "9. Governance", 1)
    add_labeled_fields(doc, [
        ("Policy owner", "NAME/TITLE"), ("Privacy contact", "EMAIL/PHONE"),
        ("Approved by", "BOARD/OFFICER"), ("Effective date", "DATE"), ("Annual review month", "MONTH"),
    ])
    return save(doc, "Data-Privacy-and-Record-Retention-Policy.docx")


def build_insurance_checklist():
    doc = new_document("Insurance Requirements and Coverage Checklist", "Planning targets for organizational, host, and participant protection", "Risk checklist", compact=True)
    add_legal_review_note(doc)
    add_callout(doc, "Broker confirmation required", "Coverage names, limits, exclusions, territory, and additional-insured requirements vary. These are planning targets only. A qualified broker should confirm the final program design, countries, activities, vehicles, housing, vulnerable populations, and contracts.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "1. Hope Sojourns organizational coverage", 1)
    add_table(doc, ["Coverage", "Planning target", "Confirm"], [
        ("Commercial general liability", "$1M per occurrence / $2M aggregate", "☐ Limits  ☐ Territory  ☐ Volunteers/interns  ☐ Additional insured"),
        ("Professional liability / errors and omissions", "$1M", "☐ Placement/advising activity  ☐ Defense costs  ☐ International claims"),
        ("Abuse and molestation / sexual misconduct", "$1M or broker recommendation", "☐ Participant-to-client  ☐ Leader-to-participant  ☐ Defense and reporting"),
        ("Directors and officers", "$1M or broker recommendation", "☐ Employment practices  ☐ Volunteer leadership  ☐ Defense"),
        ("Cyber/privacy liability", "$1M or broker recommendation", "☐ Breach response  ☐ Participant data  ☐ Vendor events"),
        ("Hired and non-owned automobile", "$1M when applicable", "☐ Rental vehicles  ☐ Volunteer drivers  ☐ International exclusion"),
        ("Workers' compensation", "As required", "☐ Employee classification  ☐ Travel  ☐ State/country obligations"),
        ("Foreign package / international liability", "Broker recommendation", "☐ Foreign voluntary workers comp  ☐ Evacuation  ☐ Kidnap/security assistance"),
    ], [2800, 2300, 4260], font_size=8.2)
    add_heading(doc, "2. Host ministry evidence", 1)
    add_checkbox_list(doc, [
        "Certificate or other acceptable evidence of general liability received",
        "Professional, abuse/molestation, automobile, property/housing, and workers' compensation coverage reviewed as applicable",
        "Policy territory includes the placement activities and location",
        "Hope Sojourns and university additional-insured requirements confirmed",
        "Material exclusions, deductibles, and claims-reporting procedures documented",
        "Host indemnity and insurance wording reviewed by counsel/broker",
    ])
    add_heading(doc, "3. Student travel coverage", 1)
    add_table(doc, ["Coverage", "Planning minimum", "Verified details"], [
        ("Travel medical", "$100,000 or university/broker requirement", "Carrier, policy, dates, territory, exclusions"),
        ("Emergency medical evacuation", "$250,000 or university/broker requirement", "Assistance number, approval process, destination"),
        ("Repatriation of remains", "$25,000 or university/broker requirement", "Benefit and contact"),
        ("Trip cancellation/interruption", "Strongly recommended", "Covered reasons and limits"),
        ("Personal liability", "As required by destination/host", "Territory and activities"),
        ("High-risk activities or conditions", "Specific written confirmation", "Adventure, driving, manual work, pre-existing conditions"),
    ], [2700, 3100, 3560], font_size=8.5)
    # Start the placement-specific record on a clean page with its heading.
    doc.add_page_break()
    add_heading(doc, "4. Placement insurance record", 1)
    add_labeled_fields(doc, [
        ("Student", "NAME"), ("Host/location", "DETAILS"), ("Dates", "DATES"),
        ("Hope Sojourns policies reviewed by", "NAME/DATE"), ("Host evidence reviewed by", "NAME/DATE"),
        ("Student policy reviewed by", "NAME/DATE"), ("Open exclusions or conditions", "DETAILS"),
        ("Broker/counsel sign-off", "NAME/DATE/NOTES"),
    ])
    add_heading(doc, "5. Before departure", 1)
    add_checkbox_list(doc, [
        "All required policies are active for the full placement and travel dates.",
        "24/7 assistance and claims contacts are in the emergency plan.",
        "The student knows whether providers require pre-authorization or payment up front.",
        "Medical evacuation destination and authorization process are understood.",
        "Vehicle use, drivers, housing, ministry activities, and vulnerable-population work are covered or excluded explicitly.",
        "Certificates, policy summaries, and relevant exclusions are stored securely and accessible during an emergency.",
    ])
    return save(doc, "Insurance-Requirements-and-Coverage-Checklist.docx")


def build_all():
    outputs = [
        build_business_model(),
        build_university_agreement(),
        build_host_agreement(),
        build_student_agreement(),
        build_learning_agreement(),
        build_code_of_conduct(),
        build_safeguarding_policy(),
        build_background_policy(),
        build_refund_policy(),
        build_emergency_plan(),
        build_parent_authorization(),
        build_privacy_policy(),
        build_insurance_checklist(),
    ]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    build_all()
