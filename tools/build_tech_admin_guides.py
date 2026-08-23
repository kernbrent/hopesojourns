"""Build the Hope Sojourns technical administration guides as Word documents.

The Markdown files in ``docs/tech-admin`` remain the maintainable source. This
script applies one shared, branded Word design system to both guides so future
revisions can be regenerated without hand-formatting the documents.

Document design: compact_reference_guide preset with named Hope Sojourns brand
overrides for typography, color, the editorial cover, and palette swatches.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs" / "tech-admin"
LOGO_PATH = PROJECT_ROOT / "assets" / "hope-sojourns-logo.png"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

# Hope Sojourns brand overrides to the compact_reference_guide preset.
COLORS = {
    "ink": "19322B",
    "muted": "60726B",
    "paper": "FFFDF8",
    "white": "FFFFFF",
    "cream": "F6F1E7",
    "forest": "275D4D",
    "forest_dark": "173F35",
    "forest_deep": "0B2720",
    "gold": "D99B42",
    "gold_light": "F3C780",
    "coral": "C9674E",
    "coral_dark": "A34C38",
    "forest_wash": "EDF5F1",
    "forest_soft": "D8EBE4",
    "sage": "73A997",
    "gold_wash": "FFF8E9",
    "gold_soft": "F4E5C8",
    "gold_ink": "6B4B1E",
    "coral_wash": "FFF0EA",
    "coral_soft": "F4D1C3",
    "success": "1F6849",
    "success_wash": "E2F1E8",
    "info": "285B8F",
    "info_wash": "E5EEF9",
    "warning": "7B4F09",
    "error": "8F3030",
}

BODY_FONT = "Segoe UI"
HEADING_FONT = "Georgia"
CODE_FONT = "Consolas"


@dataclass(frozen=True)
class GuideSpec:
    source: str
    output: str
    title: str
    subtitle: str
    running_label: str


GUIDES = (
    GuideSpec(
        source="Hope-Sojourns-Style-Guide.md",
        output="Hope-Sojourns-Style-Guide.docx",
        title="Hope Sojourns\nWebsite Style Guide",
        subtitle="Brand, color, typography, components, imagery, voice, and accessibility",
        running_label="Website Style Guide",
    ),
    GuideSpec(
        source="Hope-Sojourns-Developer-Guide.md",
        output="Hope-Sojourns-Developer-Guide.docx",
        title="Hope Sojourns\nDeveloper Guide",
        subtitle="Architecture, workflows, integrations, validation, security, and maintenance",
        running_label="Developer Guide",
    ),
)


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D8EBE4", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_paragraph_bottom_border(paragraph, color: str, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, result, end])
    set_run_font(run, size=8.5, color=COLORS["muted"])


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_settings = {
        "Heading 1": (HEADING_FONT, 16, COLORS["forest_dark"], 18, 10),
        "Heading 2": (HEADING_FONT, 13, COLORS["forest"], 14, 7),
        "Heading 3": (BODY_FONT, 12, COLORS["forest_dark"], 10, 5),
    }
    for style_name, (font_name, size, color, before, after) in heading_settings.items():
        style = styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.page_break_before = False

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(COLORS["ink"])
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code_style = styles.add_style("Hope Code", 1)
    code_style.font.name = CODE_FONT
    code_style._element.rPr.rFonts.set(qn("w:ascii"), CODE_FONT)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), CODE_FONT)
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = rgb(COLORS["forest_deep"])
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.keep_together = True

    note_style = styles.add_style("Hope Note", 1)
    note_style.font.name = BODY_FONT
    note_style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    note_style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    note_style.font.size = Pt(10.5)
    note_style.font.color.rgb = rgb(COLORS["forest_dark"])
    note_style.paragraph_format.left_indent = Inches(0.18)
    note_style.paragraph_format.right_indent = Inches(0.12)
    note_style.paragraph_format.space_before = Pt(5)
    note_style.paragraph_format.space_after = Pt(8)
    note_style.paragraph_format.line_spacing = 1.2


def create_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_num = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        if kind == "number":
            marker = "%1." if level == 0 else ("%1.%2." if level == 1 else "%1.%2.%3.")
        elif kind == "checkbox":
            marker = "☐"
        else:
            marker = "•" if level == 0 else ("–" if level == 1 else "◦")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), BODY_FONT)
        fonts.set(qn("w:hAnsi"), BODY_FONT)
        r_pr.append(fonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), COLORS["forest"])
        r_pr.append(color)
        lvl.append(r_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id_node = OxmlElement("w:abstractNumId")
    abstract_id_node.set(qn("w:val"), str(abstract_id))
    num.append(abstract_id_node)
    if kind == "number":
        for level in range(3):
            level_override = OxmlElement("w:lvlOverride")
            level_override.set(qn("w:ilvl"), str(level))
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            level_override.append(start_override)
            num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    paragraph.paragraph_format.left_indent = Inches(0.375 + min(level, 2) * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[.+?\]\(.+?\))")


def add_inline_text(paragraph, text: str, *, base_size: float = 11, base_color: str | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=base_size, color=base_color or COLORS["ink"])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, color=base_color or COLORS["ink"], bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=CODE_FONT, size=max(8.5, base_size - 1), color=COLORS["forest_dark"])
            r_pr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), COLORS["forest_wash"])
            r_pr.append(shd)
        else:
            link_match = re.match(r"\[(.+?)\]\((.+?)\)", token)
            if link_match:
                label, target = link_match.groups()
                part = paragraph.part
                rel_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), rel_id)
                link_run = OxmlElement("w:r")
                r_pr = OxmlElement("w:rPr")
                color = OxmlElement("w:color")
                color.set(qn("w:val"), COLORS["forest"])
                r_pr.append(color)
                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                r_pr.append(underline)
                link_run.append(r_pr)
                text_node = OxmlElement("w:t")
                text_node.text = label
                link_run.append(text_node)
                hyperlink.append(link_run)
                paragraph._p.append(hyperlink)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=base_size, color=base_color or COLORS["ink"])


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Hope Code")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COLORS["cream"])
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    for edge, edge_color in (("left", COLORS["gold"]), ("top", COLORS["forest_soft"]), ("bottom", COLORS["forest_soft"]), ("right", COLORS["forest_soft"])):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "14" if edge == "left" else "4")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), edge_color)
        p_bdr.append(border)
    p_pr.append(p_bdr)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, name=CODE_FONT, size=8.5, color=COLORS["forest_deep"])
        if index < len(lines) - 1:
            run.add_break()


def choose_table_widths(headers: list[str], rows: list[list[str]]) -> list[int]:
    count = len(headers)
    lower = [header.lower() for header in headers]
    if count == 2:
        if lower[0] in {"date", "change", "asset", "source", "path", "foreground on background"}:
            return [2700, 6660]
        return [2400, 6960]
    if count == 3:
        if lower == ["route", "purpose", "additional runtime"]:
            return [1650, 5040, 2670]
        return [2200, 3580, 3580]
    if count == 4:
        return [2100, 2500, 1580, 3180]
    if count == 5 and "hex" in lower and "rgb" in lower:
        return [1820, 1500, 1200, 1840, 3000]

    lengths = []
    all_rows = [headers, *rows]
    for col in range(count):
        max_len = max(len(row[col]) if col < len(row) else 0 for row in all_rows)
        lengths.append(max(8, min(max_len, 45)))
    raw = [int(CONTENT_WIDTH_DXA * length / sum(lengths)) for length in lengths]
    raw[-1] += CONTENT_WIDTH_DXA - sum(raw)
    return raw


def text_contrast(hex_value: str) -> str:
    value = hex_value.lstrip("#")
    r, g, b = int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)
    luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
    return COLORS["forest_deep"] if luminance > 0.62 else COLORS["white"]


def add_markdown_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    widths = choose_table_widths(headers, rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_row_repeat_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLORS["forest_wash"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        add_inline_text(paragraph, header, base_size=9, base_color=COLORS["forest_dark"])
        for run in paragraph.runs:
            run.bold = True

    hex_col = next((i for i, value in enumerate(headers) if value.strip().lower() == "hex"), None)
    for row_values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(1.5)
            paragraph.paragraph_format.line_spacing = 1.0
            cell_color = COLORS["ink"]
            if hex_col is not None and index == hex_col:
                match = re.search(r"#([0-9A-Fa-f]{6})", value)
                if match:
                    fill = match.group(1).upper()
                    set_cell_shading(cell, fill)
                    cell_color = text_contrast(fill)
            add_inline_text(paragraph, value, base_size=8.4, base_color=cell_color)
            if hex_col is not None and index == hex_col:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_line(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def extract_metadata(lines: list[str]) -> tuple[str, str]:
    version = next((line.split("Version", 1)[1].strip() for line in lines if line.startswith("Version ")), "1.0")
    reviewed = next((line.split(":", 1)[1].strip() for line in lines if line.startswith("Last reviewed:")), "")
    return version, reviewed


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_running_furniture(section, running_label: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run(f"HOPE SOJOURNS  |  {running_label.upper()}")
    set_run_font(left, size=8, color=COLORS["muted"], bold=True)
    tab = paragraph.add_run("\t")
    set_run_font(tab, size=8, color=COLORS["muted"])
    right = paragraph.add_run("CHRISTIAN STEPS MINISTRIES")
    set_run_font(right, size=8, color=COLORS["muted"])
    set_paragraph_bottom_border(paragraph, COLORS["gold"], "6")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(3)
    label = paragraph.add_run("Hope Sojourns technical administration  •  ")
    set_run_font(label, size=8.5, color=COLORS["muted"])
    add_page_field(paragraph)


def add_cover(doc: Document, spec: GuideSpec, version: str, reviewed: str) -> None:
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(28)
    logo_run = logo_p.add_run()
    logo_run.add_picture(str(LOGO_PATH), width=Inches(5.55))
    drawing = logo_run._r.find(qn("w:drawing"))
    if drawing is not None:
        doc_pr = drawing.find(".//" + qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", "Hope Sojourns logo with a sunrise, path, hills, and the words Go with hope. Serve with faith.")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("TECHNICAL ADMINISTRATION")
    set_run_font(run, size=10, color=COLORS["gold_ink"], bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    for index, line in enumerate(spec.title.split("\n")):
        run = title.add_run(line)
        set_run_font(run, name=HEADING_FONT, size=28, color=COLORS["forest_deep"], bold=True)
        if index < len(spec.title.split("\n")) - 1:
            run.add_break()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.left_indent = Inches(0.35)
    subtitle.paragraph_format.right_indent = Inches(0.35)
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run(spec.subtitle)
    set_run_font(run, size=12.5, color=COLORS["forest"], italic=True)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.left_indent = Inches(1.35)
    rule.paragraph_format.right_indent = Inches(1.35)
    rule.paragraph_format.space_after = Pt(22)
    set_paragraph_bottom_border(rule, COLORS["gold"], "12")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    first = meta.add_run(f"Version {version}")
    set_run_font(first, size=10, color=COLORS["muted"], bold=True)
    sep = meta.add_run("  •  ")
    set_run_font(sep, size=10, color=COLORS["gold"])
    second = meta.add_run(f"Last reviewed {reviewed}")
    set_run_font(second, size=10, color=COLORS["muted"])

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org.paragraph_format.space_before = Pt(6)
    run = org.add_run("Hope Sojourns  |  Christian Steps Ministries")
    set_run_font(run, size=9.5, color=COLORS["muted"])
    org.add_run().add_break(WD_BREAK.PAGE)


def add_contents_page(doc: Document, headings: list[str], numbering_ids: dict[str, int]) -> None:
    heading = doc.add_paragraph("Guide map", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph()
    add_inline_text(intro, "Use Word’s Navigation Pane or the section list below to move through this living reference guide.")
    for item in headings:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, numbering_ids["contents"], 0)
        display = re.sub(r"^\d+\.\s*", "", item)
        add_inline_text(paragraph, display, base_size=10.5)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_note_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Hope Note")
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COLORS["gold_wash"])
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), COLORS["gold"])
    p_bdr.append(left)
    p_pr.append(p_bdr)
    add_inline_text(paragraph, text, base_size=10.5, base_color=COLORS["forest_dark"])


def preserve_list_boundaries(paragraph, lines: list[str], index: int, pattern: str) -> None:
    """Avoid leaving the first or final list item isolated across a page break."""
    previous_is_list = index > 0 and re.match(pattern, lines[index - 1]) is not None
    next_is_list = index + 1 < len(lines) and re.match(pattern, lines[index + 1]) is not None
    following_is_list = index + 2 < len(lines) and re.match(pattern, lines[index + 2]) is not None
    paragraph.paragraph_format.keep_together = True
    if next_is_list and (not previous_is_list or not following_is_list):
        paragraph.paragraph_format.keep_with_next = True


def render_markdown(doc: Document, lines: list[str], numbering_ids: dict[str, int]) -> None:
    index = 0
    in_code = False
    code_lines: list[str] = []
    active_number_id: int | None = None

    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not line.strip():
            index += 1
            continue

        if line.startswith("| ") and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            active_number_id = None
            headers = parse_table_line(line)
            index += 2
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                row = parse_table_line(lines[index])
                if len(row) == len(headers):
                    rows.append(row)
                index += 1
            add_markdown_table(doc, headers, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            active_number_id = None
            markdown_level = len(heading_match.group(1))
            word_level = min(markdown_level - 1, 3)
            paragraph = doc.add_paragraph(style=f"Heading {word_level}")
            add_inline_text(paragraph, heading_match.group(2), base_size={1: 16, 2: 13, 3: 12}[word_level], base_color=COLORS["forest_dark"] if word_level != 2 else COLORS["forest"])
            for run in paragraph.runs:
                run.bold = True
                if word_level in (1, 2):
                    set_run_font(run, name=HEADING_FONT, size={1: 16, 2: 13}[word_level], color=COLORS["forest_dark"] if word_level == 1 else COLORS["forest"], bold=True)
            index += 1
            continue

        checkbox = re.match(r"^(\s*)-\s+\[([ xX])\]\s+(.+)$", line)
        if checkbox:
            active_number_id = None
            level = min(len(checkbox.group(1)) // 2, 2)
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, numbering_ids["checkbox"], level)
            preserve_list_boundaries(paragraph, lines, index, r"^(\s*)-\s+\[([ xX])\]\s+(.+)$")
            add_inline_text(paragraph, checkbox.group(3))
            index += 1
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet:
            active_number_id = None
            level = min(len(bullet.group(1)) // 2, 2)
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, numbering_ids["bullet"], level)
            preserve_list_boundaries(paragraph, lines, index, r"^(\s*)[-*]\s+(.+)$")
            add_inline_text(paragraph, bullet.group(2))
            index += 1
            continue

        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered:
            if active_number_id is None:
                active_number_id = create_numbering(doc, "number")
            level = min(len(numbered.group(1)) // 2, 2)
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, active_number_id, level)
            preserve_list_boundaries(paragraph, lines, index, r"^(\s*)\d+\.\s+(.+)$")
            add_inline_text(paragraph, numbered.group(2))
            index += 1
            continue

        if line.startswith("> "):
            active_number_id = None
            add_note_paragraph(doc, line[2:])
            index += 1
            continue

        active_number_id = None
        paragraph = doc.add_paragraph()
        add_inline_text(paragraph, line)
        index += 1


def mark_document_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build_guide(spec: GuideSpec) -> Path:
    source_path = DOCS_DIR / spec.source
    output_path = DOCS_DIR / spec.output
    lines = source_path.read_text(encoding="utf-8").splitlines()
    version, reviewed = extract_metadata(lines)
    section_headings = [match.group(1) for line in lines if (match := re.match(r"^##\s+(.+)$", line))]

    doc = Document()
    section = doc.sections[0]
    configure_page(section)
    configure_styles(doc)
    add_running_furniture(section, spec.running_label)
    mark_document_update_fields(doc)

    doc.core_properties.title = spec.title.replace("\n", " ")
    doc.core_properties.subject = spec.subtitle
    doc.core_properties.author = "Christian Steps Ministries"
    doc.core_properties.last_modified_by = "Christian Steps Ministries"
    doc.core_properties.keywords = "Hope Sojourns, technical administration, style guide, developer guide"
    doc.core_properties.comments = "Generated from the canonical Markdown source in docs/tech-admin."

    numbering_ids = {
        "bullet": create_numbering(doc, "bullet"),
        "number": create_numbering(doc, "number"),
        "checkbox": create_numbering(doc, "checkbox"),
        "contents": create_numbering(doc, "bullet"),
    }

    add_cover(doc, spec, version, reviewed)
    add_contents_page(doc, section_headings, numbering_ids)

    start_index = 1
    while start_index < len(lines) and (not lines[start_index].strip() or lines[start_index].startswith("Version ") or lines[start_index].startswith("Last reviewed:")):
        start_index += 1
    render_markdown(doc, lines[start_index:], numbering_ids)

    # Keep the final paragraph compact and avoid an accidental blank last page.
    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=DOCS_DIR, help="Destination for generated Word documents")
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    for spec in GUIDES:
        path = build_guide(spec)
        if args.output_dir.resolve() != DOCS_DIR.resolve():
            destination = args.output_dir / path.name
            destination.write_bytes(path.read_bytes())
            path = destination
        print(path)


if __name__ == "__main__":
    main()
