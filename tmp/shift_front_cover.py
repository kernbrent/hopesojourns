from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\kernb\OneDrive\MasterFolder\Documents\HopeSojourns")
DOCX = ROOT / "brochure" / "Hope-Sojourns-Trifold-Brochure-Full-Page.docx"

# A 0.65 cm shift preserves a safe outside print edge while correcting the
# folded cover alignment shown in the physical proof.
SHIFT_DXA = 369


def set_indent(paragraph, tag, value):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn(f"w:{tag}"), str(value))


doc = Document(DOCX)
front_cover = doc.tables[0].rows[0].cells[2]

for paragraph in front_cover.paragraphs:
    set_indent(paragraph, "left", SHIFT_DXA)
    set_indent(paragraph, "right", -SHIFT_DXA)

doc.core_properties.comments = (
    "Full-page, exact-third trifold version. Front-cover composition shifted "
    "0.65 cm toward the outside right edge for improved alignment after folding."
)
doc.save(DOCX)
print(DOCX)
