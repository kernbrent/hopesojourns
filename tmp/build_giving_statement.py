from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\kernb\OneDrive\MasterFolder\Documents\HopeSojourns")
OUT = ROOT / "admin" / "supplemental-documents" / "Hope-Sojourns-Giving-Statement-Template.docx"
LOGO = ROOT / "assets" / "hope-sojourns-logo.png"

INK = "19322B"
FOREST = "275D4D"
FOREST_DARK = "173F35"
GOLD = "D99B42"
CREAM = "F6F1E7"
PALE = "FFF8E9"
MUTED = "60726B"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def borders(table, color="D7DED9", size="6"):
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    block = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        block.append(edge)
    tbl_pr.append(block)


def font(run, size=10.5, color=INK, bold=False, name="Aptos"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def add_para(container, text="", *, size=10.5, color=INK, bold=False, before=0, after=7, align=None, name="Aptos"):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    font(r, size=size, color=color, bold=bold, name=name)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(.55)
section.bottom_margin = Inches(.55)
section.left_margin = Inches(.7)
section.right_margin = Inches(.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.08

header = doc.add_table(rows=1, cols=2)
header.alignment = WD_TABLE_ALIGNMENT.CENTER
header.autofit = False
header.columns[0].width = Inches(3.5)
header.columns[1].width = Inches(3.5)
header.cell(0, 0).width = Inches(3.5)
header.cell(0, 1).width = Inches(3.5)
for cell in header.rows[0].cells:
    set_cell_margins(cell, 0, 0, 0, 0)
logo_p = header.cell(0, 0).paragraphs[0]
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
shape = logo_p.add_run().add_picture(str(LOGO), width=Inches(2.65))
shape._inline.docPr.set("descr", "Hope Sojourns logo")
meta_p = header.cell(0, 1).paragraphs[0]
meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for text, size, color, bold in (
    ("OFFICIAL CONTRIBUTION RECEIPT", 8.5, GOLD, True),
    ("[[TAX_YEAR]] GIVING STATEMENT", 15, FOREST_DARK, True),
    ("HopeSojourns.com", 9, MUTED, True),
):
    r = meta_p.add_run(text + "\n")
    font(r, size=size, color=color, bold=bold, name="Georgia" if size == 15 else "Aptos")

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(3)
rule.paragraph_format.space_after = Pt(10)
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "16")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), GOLD)
p_bdr.append(bottom)
p_pr.append(p_bdr)

address = doc.add_table(rows=1, cols=2)
address.alignment = WD_TABLE_ALIGNMENT.CENTER
address.autofit = False
address.columns[0].width = Inches(4.35)
address.columns[1].width = Inches(2.65)
left, right = address.rows[0].cells
for cell in (left, right):
    shade(cell, CREAM)
    set_cell_margins(cell, 130, 160, 130, 160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
left_p = left.paragraphs[0]
left_p.paragraph_format.space_after = Pt(0)
for line in ("[[DONOR_NAME]]", "[[ADDRESS_LINE_1]]", "[[ADDRESS_LINE_2_OPTIONAL]]", "[[CITY_STATE_ZIP]]"):
    r = left_p.add_run(line + "\n")
    font(r, size=9.5, color=INK, bold=line == "[[DONOR_NAME]]")
right_p = right.paragraphs[0]
right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
right_p.paragraph_format.space_after = Pt(0)
for line in ("Letter date", "[[LETTER_DATE]]", "Receipt number", "[[RECEIPT_NUMBER]]"):
    r = right_p.add_run(line + "\n")
    font(r, size=8.8, color=MUTED if "[[" not in line else INK, bold="[[" in line)

add_para(doc, "Dear [[GREETING_NAME]],", before=10, after=7)
add_para(
    doc,
    "Thank you for your generous support of Hope Sojourns. Your partnership helps prepare people to serve, strengthen trusted ministry relationships, create training resources, and respond responsibly to needs identified by local leaders.",
)
add_para(
    doc,
    "Christian Steps Ministries receives and processes charitable gifts made in support of Hope Sojourns. Christian Steps Ministries is a tax-exempt organization recognized under Section 501(c)(3) of the Internal Revenue Code (EIN 81-1678503).",
    after=10,
)

summary = doc.add_table(rows=1, cols=2)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
summary.autofit = False
summary.columns[0].width = Inches(4.4)
summary.columns[1].width = Inches(2.6)
for cell in summary.rows[0].cells:
    shade(cell, FOREST_DARK)
    set_cell_margins(cell, 150, 180, 150, 180)
summary.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
summary.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = summary.cell(0, 0).paragraphs[0].add_run("TOTAL CHARITABLE GIVING")
font(r, size=9, color=WHITE, bold=True)
r = summary.cell(0, 1).paragraphs[0].add_run("[[TOTAL_GIFT_AMOUNT]]")
font(r, size=14, color=WHITE, bold=True, name="Georgia")

add_para(doc, "CONTRIBUTION DETAILS", size=8.5, color=GOLD, bold=True, before=12, after=5)
details = doc.add_table(rows=2, cols=4)
details.alignment = WD_TABLE_ALIGNMENT.CENTER
details.autofit = False
widths = (1.22, 1.22, 2.88, 1.68)
for i, width in enumerate(widths):
    details.columns[i].width = Inches(width)
for cell in details.rows[0].cells:
    shade(cell, FOREST)
    set_cell_margins(cell, 100, 110, 100, 110)
for cell, label in zip(details.rows[0].cells, ("Date", "Amount", "Designation", "Payment method")):
    r = cell.paragraphs[0].add_run(label)
    font(r, size=8.5, color=WHITE, bold=True)
for cell, value in zip(details.rows[1].cells, ("[[GIFT_DATE]]", "[[GIFT_AMOUNT]]", "[[DESIGNATION]]", "[[PAYMENT_METHOD]]")):
    set_cell_margins(cell, 120, 110, 120, 110)
    r = cell.paragraphs[0].add_run(value)
    font(r, size=8.7, color=INK)
borders(details)

note = doc.add_table(rows=1, cols=1)
note.alignment = WD_TABLE_ALIGNMENT.CENTER
shade(note.cell(0, 0), PALE)
set_cell_margins(note.cell(0, 0), 130, 160, 130, 160)
note_p = note.cell(0, 0).paragraphs[0]
note_p.paragraph_format.space_after = Pt(0)
r = note_p.add_run("Tax statement: ")
font(r, size=9.2, color=INK, bold=True)
r = note_p.add_run("No goods or services were provided in exchange for these charitable contributions. Please retain this statement with your tax records.")
font(r, size=9.2, color=INK)

add_para(doc, "With gratitude,", before=12, after=2)
add_para(doc, "Hope Sojourns", size=13, color=FOREST_DARK, bold=True, after=0, name="Georgia")
add_para(doc, "in partnership with Christian Steps Ministries", size=9.2, color=MUTED, after=0)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(4)
footer_p.paragraph_format.space_after = Pt(0)
r = footer_p.add_run("GO WITH HOPE. SERVE WITH FAITH.  |  HopeSojourns.com/giving/")
font(r, size=8, color=FOREST_DARK, bold=True)

doc.core_properties.title = "Hope Sojourns Giving Statement Template"
doc.core_properties.subject = "Annual donor contribution receipt template"
doc.core_properties.author = "Hope Sojourns"
doc.core_properties.comments = "Editable donor giving statement template. Replace bracketed placeholders before issuing."
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
