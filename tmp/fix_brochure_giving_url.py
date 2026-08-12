from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from reportlab.graphics.barcode import qr


ROOT = Path(r"C:\Users\kernb\OneDrive\MasterFolder\Documents\HopeSojourns")
DOCX = ROOT / "brochure" / "Hope-Sojourns-Trifold-Brochure-Full-Page.docx"
QR_PNG = ROOT / "assets" / "hope-sojourns-giving-qr.png"
URL = "https://hopesojourns.com/giving/"
DISPLAY_URL = "HopeSojourns.com/giving/"


def make_qr():
    widget = qr.QrCodeWidget(URL, barLevel="H")
    widget.draw()
    modules = widget.qr.modules
    module_count = len(modules)
    quiet_zone = 4
    module_pixels = 12
    qr_pixels = (module_count + 2 * quiet_zone) * module_pixels
    qr_image = Image.new("RGB", (qr_pixels, qr_pixels), "white")
    qr_draw = ImageDraw.Draw(qr_image)
    for row, values in enumerate(modules):
        for column, value in enumerate(values):
            if value:
                x0 = (column + quiet_zone) * module_pixels
                y0 = (row + quiet_zone) * module_pixels
                qr_draw.rectangle(
                    (x0, y0, x0 + module_pixels - 1, y0 + module_pixels - 1),
                    fill="black",
                )

    canvas = Image.new("RGB", (498, 550), "white")
    target_size = 468
    qr_image = qr_image.resize((target_size, target_size), Image.Resampling.NEAREST)
    canvas.paste(qr_image, ((canvas.width - target_size) // 2, 67))
    draw = ImageDraw.Draw(canvas)
    try:
        font = ImageFont.truetype("arial.ttf", 27)
    except OSError:
        font = ImageFont.load_default()
    title = "Hope Sojourns"
    bbox = draw.textbbox((0, 0), title, font=font)
    draw.text(((canvas.width - (bbox[2] - bbox[0])) / 2, 12), title, fill="black", font=font)
    canvas.save(QR_PNG)


def replace_hyperlink_text_and_target(doc):
    back_cover = doc.tables[0].rows[0].cells[1]
    link_paragraph = next(p for p in back_cover.paragraphs if "HopeSojourns.com/give/" in p.text)
    hyperlink = link_paragraph._p.find(qn("w:hyperlink"))
    rid = hyperlink.get(qn("r:id"))
    doc.part.rels[rid]._target = URL
    text_node = hyperlink.find(".//" + qn("w:t"))
    text_node.text = DISPLAY_URL


def replace_qr_image(doc):
    qr_shape = next(
        shape for shape in doc.inline_shapes
        if shape._inline.docPr.get("name") == "Hope Sojourns donation QR code"
    )
    blip = qr_shape._inline.find(".//" + qn("a:blip"))
    old_rid = blip.get(qn("r:embed"))
    image_part = doc.part.related_parts[old_rid]
    image_part._blob = QR_PNG.read_bytes()
    qr_shape._inline.docPr.set("descr", "QR code linking to the Hope Sojourns giving page")


make_qr()
doc = Document(DOCX)
replace_hyperlink_text_and_target(doc)
replace_qr_image(doc)
doc.core_properties.comments = (
    "Full-page exact-third trifold. Front cover shifted toward the outside edge; "
    "giving URL and QR code corrected to hopesojourns.com/giving/."
)
doc.save(DOCX)
print(DOCX)
print(QR_PNG)
