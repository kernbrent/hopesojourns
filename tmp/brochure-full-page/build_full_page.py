from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\kernb\OneDrive\MasterFolder\Documents\HopeSojourns")
SOURCE = ROOT / "brochure" / "Hope-Sojourns-Trifold-Brochure.docx"
OUTPUT = ROOT / "tmp" / "brochure-full-page" / "Hope-Sojourns-Trifold-Brochure-Full-Page.docx"
QR = ROOT / "tmp" / "brochure-full-page" / "hope-sojourns-qr.png"
WITLOOL = ROOT / "assets" / "witlool-mark-transparent.png"

INK = "19322B"
MUTED = "60726B"
FOREST_DARK = "173F35"
GOLD = "D99B42"
PALE_GOLD = "FBF1DE"


def set_font(run, name: str, size: float, color: str, *, bold: bool = False, italic: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), name)


def set_paragraph_fill(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color: str = GOLD):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 8.4):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Segoe UI")
    fonts.set(qn("w:hAnsi"), "Segoe UI")
    r_pr.append(fonts)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), FOREST_DARK)
    r_pr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(round(size * 2)))
    r_pr.append(sz)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_full_page_table(table):
    tbl_pr = table._tbl.tblPr

    width = tbl_pr.find(qn("w:tblW"))
    width.set(qn("w:w"), "15840")
    width.set(qn("w:type"), "dxa")

    justification = tbl_pr.find(qn("w:jc"))
    if justification is None:
        justification = OxmlElement("w:jc")
        tbl_pr.append(justification)
    justification.set(qn("w:val"), "left")

    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tbl_pr.append(borders)

    grid_widths = (5280, 5280, 5280)
    for grid_col, width_value in zip(table._tbl.tblGrid.gridCol_lst, grid_widths):
        grid_col.set(qn("w:w"), str(width_value))

    for cell, width_value in zip(table.rows[0].cells, grid_widths):
        tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        tc_width.set(qn("w:w"), str(width_value))
        tc_width.set(qn("w:type"), "dxa")
        set_cell_margins(cell, top=270, start=245, bottom=270, end=245)

    row_pr = table.rows[0]._tr.get_or_add_trPr()
    for height in list(row_pr.findall(qn("w:trHeight"))):
        row_pr.remove(height)
    row_height = OxmlElement("w:trHeight")
    row_height.set(qn("w:val"), "11300")
    row_height.set(qn("w:hRule"), "atLeast")
    row_pr.append(row_height)


def configure_bottom_mark_table(table):
    tbl_pr = table._tbl.tblPr

    width = tbl_pr.find(qn("w:tblW"))
    width.set(qn("w:w"), "4680")
    width.set(qn("w:type"), "dxa")

    justification = tbl_pr.find(qn("w:jc"))
    if justification is None:
        justification = OxmlElement("w:jc")
        tbl_pr.append(justification)
    justification.set(qn("w:val"), "center")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    tbl_pr.append(borders)

    for grid_col in table._tbl.tblGrid.gridCol_lst:
        grid_col.set(qn("w:w"), "2340")

    for cell in table.rows[0].cells:
        tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        tc_width.set(qn("w:w"), "2340")
        tc_width.set(qn("w:type"), "dxa")
        set_cell_margins(cell, top=0, start=40, bottom=0, end=40)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


doc = Document(SOURCE)
section = doc.sections[0]
section.left_margin = Inches(0)
section.right_margin = Inches(0)
section.top_margin = Inches(0)
section.bottom_margin = Inches(0)
section.header_distance = Inches(0)
section.footer_distance = Inches(0)

for table in doc.tables:
    configure_full_page_table(table)

# Replace the back-cover Athens outreach image with the user-provided QR code.
back_cover = doc.tables[0].rows[0].cells[1]
photo_paragraph = back_cover.paragraphs[12]
remove_paragraph(photo_paragraph)

donate_eyebrow = back_cover.add_paragraph()
donate_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
donate_eyebrow.paragraph_format.space_before = Pt(6)
donate_eyebrow.paragraph_format.space_after = Pt(1)
donate_eyebrow.paragraph_format.keep_with_next = True
set_paragraph_fill(donate_eyebrow, PALE_GOLD)
set_paragraph_left_border(donate_eyebrow)
run = donate_eyebrow.add_run("SUPPORT THE JOURNEY")
set_font(run, "Segoe UI", 7.0, GOLD, bold=True)

donate_text = back_cover.add_paragraph()
donate_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
donate_text.paragraph_format.space_before = Pt(0)
donate_text.paragraph_format.space_after = Pt(3)
donate_text.paragraph_format.keep_with_next = True
set_paragraph_fill(donate_text, PALE_GOLD)
set_paragraph_left_border(donate_text)
run = donate_text.add_run(
    "Donate to Hope Sojourns and help equip purposeful, partner-led journeys of faith, service, and lasting relationship."
)
set_font(run, "Segoe UI", 7.5, INK)

donate_link = back_cover.add_paragraph()
donate_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
donate_link.paragraph_format.space_before = Pt(1)
donate_link.paragraph_format.space_after = Pt(1)
add_hyperlink(donate_link, "HopeSojourns.com/give/", "https://hopesojourns.com/give/", size=8.2)

bottom_marks = back_cover.add_table(rows=1, cols=2)
bottom_marks.autofit = False
configure_bottom_mark_table(bottom_marks)

qr_cell, witlool_cell = bottom_marks.rows[0].cells
qr_paragraph = qr_cell.paragraphs[0]
qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
qr_paragraph.paragraph_format.space_before = Pt(0)
qr_paragraph.paragraph_format.space_after = Pt(0)
qr_shape = qr_paragraph.add_run().add_picture(str(QR), width=Inches(1.39))
qr_shape._inline.docPr.set("name", "Hope Sojourns donation QR code")
qr_shape._inline.docPr.set("descr", "QR code for supporting Hope Sojourns")

witlool_paragraph = witlool_cell.paragraphs[0]
witlool_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
witlool_paragraph.paragraph_format.space_before = Pt(40)
witlool_paragraph.paragraph_format.space_after = Pt(0)
witlool_shape = witlool_paragraph.add_run().add_picture(str(WITLOOL), width=Inches(1.18))
witlool_shape.width = Inches(1.18)
witlool_shape.height = Inches(0.38)
blip_fill = witlool_shape._inline.xpath(".//pic:blipFill")[0]
source_rect = OxmlElement("a:srcRect")
source_rect.set("l", "20500")
source_rect.set("t", "33000")
source_rect.set("r", "19000")
source_rect.set("b", "38500")
blip_fill.insert(1, source_rect)
witlool_shape._inline.docPr.set("name", "WitLooL ministry mark")
witlool_shape._inline.docPr.set("descr", "WitLooL: Working in the Light of our Lord")

# The section properties are already a direct body child, so the empty final
# paragraph can be removed without affecting the two brochure sides.
trailing = doc.paragraphs[-1]
remove_paragraph(trailing)

doc.core_properties.title = "Hope Sojourns Trifold Brochure - Full Page"
doc.core_properties.comments = "Full-page, exact-third trifold version with donation QR code and WitLooL mark on the back cover."
doc.save(OUTPUT)
print(OUTPUT)
